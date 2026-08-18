#!/usr/bin/env python3
"""Gera o manuscrito em .docx para REVISAO dos coautores.

O .docx nao e o artefato de submissao, e o artefato de revisao. Isso decide tudo:
figura embutida no ponto do texto que fala dela (o revisor precisa ver o grafico junto
do argumento), tabela nativa do Word (imagem de tabela mata comentario por celula e
controle de alteracoes), e numeracao de linha ligada (e o que o revisor usa para dizer
"linha 214").

De onde vem cada coisa, e por que:

- PROSA: de paper/manuscript.tex, parseada. Nao reescrita aqui. A prosa ja foi
  conferida numero a numero contra o verified_numbers.json; reescreve-la neste arquivo
  criaria uma segunda fonte da verdade que diverge em silencio na proxima rodada.
- TABELAS: dos paper/tables/*.tex ja gerados, convertidos para tabela nativa. Parsear
  o .tex gerado, em vez de remontar do JSON, garante que a celula do Word e a mesma
  celula do LaTeX, com a mesma formatacao.
- FIGURAS: os PNG de 300 dpi de paper/figures/, que o build_paper_assets.py rasteriza
  a partir do mesmo pgfplots que entra no .tex.

Nao converter de LaTeX com pandoc: ele perde booktabs, estraga figura vetorial e cria
um artefato derivado que diverge do original. Aqui os dois sao irmaos do mesmo dado.

Uso:
    python scripts/build_docx.py
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
FIGURES = PAPER / "figures"
BUILD = ROOT / "build"
VERMELHO = RGBColor(0xC0, 0x00, 0x00)


# --------------------------------------------------------------------------- #
# infraestrutura de formatacao
# --------------------------------------------------------------------------- #
def setup_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # sem o eastAsia o Word troca a fonte em parte do documento
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)

    sec = doc.sections[0]
    sec.start_type = WD_SECTION_START.NEW_PAGE
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, Inches(1))
    add_line_numbers(sec)
    add_page_numbers(sec)
    return doc


def add_line_numbers(section) -> None:
    """Numeracao de linha continua. O Word nao liga por padrao e o revisor precisa."""
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    section._sectPr.append(ln)


def add_page_numbers(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        else:
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        run._r.append(el)


# --------------------------------------------------------------------------- #
# conversao de LaTeX para texto do Word
# --------------------------------------------------------------------------- #
ACENTOS = {
    r"\~a": "ã", r"\~o": "õ", r"\~A": "Ã", r"\~e": "ẽ",
    r"\'a": "á", r"\'e": "é", r"\'i": "í", r"\'o": "ó", r"\'u": "ú",
    r"\'A": "Á", r"\'E": "É", r"\'O": "Ó",
    r"\^a": "â", r"\^e": "ê", r"\^o": "ô", r"\^A": "Â", r"\^E": "Ê",
    r"\`a": "à", r"\`e": "è",
    r"\c{c}": "ç", r"\c{C}": "Ç",
}


def limpa_inline(s: str, cites: dict, refs: dict) -> list[tuple[str, dict]]:
    """Converte um trecho de LaTeX em runs (texto, formato).

    Devolve lista de (texto, {bold, italic, red}) porque a formatacao carrega
    significado: negrito e rotulo de estrutura (Background., Methods.), vermelho e
    pendencia. Achatar tudo em texto plano perderia os dois.
    """
    # 1. Comandos COM argumento que nao produzem texto. Tem que sair antes da
    #    remocao generica, senao o argumento fica orfao ("\vspace{-2em}" -> "-2em").
    for cmd in ("vspace", "hspace", "label", "footnotemark", "bibliographystyle",
                "bibliography", "graphicspath", "pagestyle", "date"):
        s = re.sub(r"\\" + cmd + r"\*?\s*(\[[^\]]*\])?\{[^{}]*\}", "", s)
    s = re.sub(r"\\thanks\{[^{}]*\}", "", s)

    # 2. Comandos SEM argumento. Precisam de fronteira explicita: colados no texto
    #    seguinte ("\noindent\textbf{X}" apos substituir -> "\noindentX"), a remocao
    #    generica levaria a palavra junto.
    for cmd in ("noindent", "maketitle", "hrule", "centering", "normalfont",
                "large", "Large", "bfseries", "itshape", "small", "footnotesize",
                "linenumbers", "clearpage", "newpage", "and"):
        s = re.sub(r"\\" + cmd + r"(?![a-zA-Z])", " ", s)
    s = s.replace("\\\\", " ")          # quebra de linha do LaTeX

    def texto_puro(x: str) -> str:
        x = re.sub(r"\\citep?\{([^}]+)\}", lambda m: "[" + ", ".join(
            str(cites.get(k.strip(), "?")) for k in m.group(1).split(",")) + "]", x)
        # \ref devolve so o numero: o texto ja escreve "Table" ou "Figure" antes dele
        x = re.sub(r"\\ref\{([^}]+)\}", lambda m: refs.get(m.group(1), "?"), x)
        x = re.sub(r"\\(url|texttt|emph|textbf|textit)\{([^}]*)\}", r"\2", x)
        for a, b in ACENTOS.items():
            x = x.replace(a, b)
        x = x.replace("{,}", ",").replace("--", "-")
        x = re.sub(r"\$([^$]*)\$", r"\1", x)
        for a, b in ((r"\%", "%"), (r"\&", "&"), (r"\_", "_"), (r"\$", "$"),
                     (r"\#", "#"), ("~", " ")):
            x = x.replace(a, b)
        x = re.sub(r"\\[a-zA-Z]+\*?", "", x)
        x = x.replace("{", "").replace("}", "")
        return re.sub(r"[ \t]+", " ", x)

    # 3. Tokeniza \missing (vermelho) e \textbf (negrito) em runs proprios
    partes: list[tuple[str, dict]] = []
    padrao = re.compile(r"\\(missing|textbf)\{")
    resto = s
    while True:
        m = padrao.search(resto)
        if not m:
            break
        antes = resto[:m.start()]
        if antes.strip():
            partes.append((texto_puro(antes), {}))
        i, prof = m.end(), 1
        while i < len(resto) and prof:
            prof += (resto[i] == "{") - (resto[i] == "}")
            i += 1
        dentro = texto_puro(resto[m.end():i - 1])
        if m.group(1) == "missing":
            partes.append(("[PENDING: " + dentro + "]", {"bold": True, "red": True}))
        else:
            partes.append((dentro, {"bold": True}))
        resto = resto[i:]
    if resto.strip():
        partes.append((texto_puro(resto), {}))
    return [x for x in partes if x[0].strip()]


def escreve_runs(p, runs):
    for texto, fmt in runs:
        r = p.add_run(texto)
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        if fmt.get("red"):
            r.font.color.rgb = VERMELHO
    return p


# --------------------------------------------------------------------------- #
# tabela booktabs -> tabela nativa
# --------------------------------------------------------------------------- #
def parse_tabela(tex: str, cites: dict, refs: dict) -> dict:
    """Le um paper/tables/*.tex gerado e devolve cabecalho, linhas, legenda e nota."""
    cap = re.search(r"\\caption\{(.*?)\}\s*\n\\label", tex, re.S)
    legenda = " ".join(x[0] for x in limpa_inline(cap.group(1), cites, refs)) if cap else ""
    corpo = re.search(r"\\midrule(.*?)\\bottomrule", tex, re.S)
    cabec = re.search(r"\\toprule(.*?)\\midrule", tex, re.S)

    def celulas(linha: str) -> list[tuple[str, bool]]:
        out = []
        for c in linha.split("&"):
            neg = "\\textbf{" in c
            txt = " ".join(x[0] for x in limpa_inline(c, cites, refs)) or ""
            out.append((txt.strip(), neg))
        return out

    header = [c for c, _ in celulas(cabec.group(1).replace("\\\\", "").strip())] if cabec else []
    linhas = []
    for l in corpo.group(1).split("\\\\"):
        if l.strip():
            linhas.append(celulas(l.strip()))
    nota = re.search(r"\\footnotesize\s*(.*?)\s*\\end\{minipage\}", tex, re.S)
    rodape = " ".join(x[0] for x in limpa_inline(nota.group(1), cites, refs)) if nota else ""
    return {"legenda": legenda, "header": header, "linhas": linhas, "nota": rodape}


def add_tabela(doc, dados, rotulo):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(rotulo)
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(" " + dados["legenda"])
    r2.font.size = Pt(10)

    t = doc.add_table(rows=1, cols=len(dados["header"]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(dados["header"]):
        cel = t.rows[0].cells[i]
        cel.text = ""
        run = cel.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        cel.paragraphs[0].paragraph_format.line_spacing = 1.0
    for linha in dados["linhas"]:
        cells = t.add_row().cells
        for ci, (txt, neg) in enumerate(linha):
            if ci >= len(cells):
                continue
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(txt)
            run.font.size = Pt(9)
            run.bold = neg
            cells[ci].paragraphs[0].paragraph_format.line_spacing = 1.0
    if dados["nota"]:
        pn = doc.add_paragraph()
        pn.paragraph_format.line_spacing = 1.0
        rn = pn.add_run(dados["nota"])
        rn.font.size = Pt(9)
        rn.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def add_figura(doc, nome, rotulo, legenda) -> bool:
    png = FIGURES / f"{nome}.png"
    if not png.exists():
        p = doc.add_paragraph()
        r = p.add_run(f"[FIGURA AUSENTE: {nome}.png]")
        r.bold = True
        r.font.color.rgb = VERMELHO
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(str(png), width=Inches(6.0))
    pc = doc.add_paragraph()
    pc.paragraph_format.line_spacing = 1.0
    pc.paragraph_format.space_after = Pt(10)
    r = pc.add_run(rotulo)
    r.bold = True
    r.font.size = Pt(10)
    r2 = pc.add_run(" " + legenda)
    r2.font.size = Pt(10)
    return True


# --------------------------------------------------------------------------- #
# indices de citacao e de referencia cruzada
# --------------------------------------------------------------------------- #
def indexa(tex: str) -> tuple[dict, dict, list[str]]:
    """Numera citacoes por ordem de aparicao (unsrtnat) e resolve \\ref."""
    cites: dict[str, int] = {}
    for m in re.finditer(r"\\citep?\{([^}]+)\}", tex):
        for k in m.group(1).split(","):
            k = k.strip()
            if k not in cites:
                cites[k] = len(cites) + 1
    refs: dict[str, str] = {}
    n_tab = n_fig = 0
    for m in re.finditer(r"\\input\{tables/(\w+)\}|\\label\{(fig:[\w]+)\}", tex):
        if m.group(1):
            n_tab += 1
            refs[f"__tab__{m.group(1)}"] = str(n_tab)
        else:
            n_fig += 1
            refs[m.group(2)] = str(n_fig)
    # label das tabelas vem de dentro do arquivo da tabela
    n_tab = 0
    for m in re.finditer(r"\\input\{tables/(\w+)\}", tex):
        n_tab += 1
        arq = (PAPER / "tables" / f"{m.group(1)}.tex").read_text(encoding="utf-8")
        lab = re.search(r"\\label\{(tab:[\w]+)\}", arq)
        if lab:
            refs[lab.group(1)] = str(n_tab)
    ordem = sorted(cites, key=cites.get)
    return cites, refs, ordem


# --------------------------------------------------------------------------- #
# bibliografia
# --------------------------------------------------------------------------- #
def limpa_bib(x: str) -> str:
    for a, b in ACENTOS.items():
        x = x.replace(a, b)
    x = x.replace("\\&", "&").replace("\\%", "%").replace("\\_", "_")
    return x.replace("{", "").replace("}", "").replace("--", "-").strip()


def formata_bib(ordem: list) -> list:
    """Referencias numeradas, na ordem de aparicao no texto (estilo unsrt)."""
    bib = (PAPER / "refs" / "references.bib").read_text(encoding="utf-8")
    entradas = {}
    for m in re.finditer(r"@\w+\{([^,]+),(.*?)\n\}", bib, re.S):
        campos = {}
        for c in re.finditer(r"(\w+)\s*=\s*\{(.*?)\}\s*(?:,|\n\})", m.group(2), re.S):
            campos[c.group(1).lower()] = re.sub(r"\s+", " ", c.group(2)).strip()
        entradas[m.group(1).strip()] = campos

    saida = []
    for chave in ordem:
        e = entradas.get(chave)
        if not e:
            saida.append(("[PENDING: entrada bib ausente para " + chave + "]", True))
            continue
        partes = [x.strip() for x in limpa_bib(e.get("author", "")).split(" and ")]
        autores = ", ".join(partes[:3]) + (", et al." if len(partes) > 3 else "")
        s = autores + ". " + limpa_bib(e.get("title", ""))
        onde = limpa_bib(e.get("journal") or e.get("booktitle") or e.get("publisher") or "")
        if onde:
            s += ". " + onde
        if e.get("year"):
            s += ". " + e["year"]
        # Sem volume, a pagina entra depois de virgula: ";:785-794" e artefato de
        # concatenar um campo vazio, e aparece em anais de congresso, que nao tem volume.
        vol = e.get("volume", "")
        pag = limpa_bib(e.get("pages", ""))
        if vol:
            s += ";" + vol + ("(" + e["number"] + ")" if e.get("number") else "")
            if pag:
                s += ":" + pag
        elif pag:
            s += ", " + pag
        if e.get("doi"):
            s += ". doi:" + e["doi"]
        elif e.get("note"):
            s += ". " + limpa_bib(e["note"])
        elif e.get("howpublished"):
            s += ". " + limpa_bib(e["howpublished"]).replace("\\url", "")
        saida.append((s.rstrip(".") + ".", False))
    return saida


# --------------------------------------------------------------------------- #
# montagem
# --------------------------------------------------------------------------- #
def carimbo(v: dict) -> str:
    """Proveniencia: de qual execucao do pipeline os numeros vieram."""
    b, d = v["backtest"], v["dados"]
    return ("Gerado de paper/verified_numbers.json. Serie {ini} a {fim} ({n} meses, "
            "{cv:,} obitos cardiovasculares). Backtesting: {jan} janelas x {h} "
            "horizontes = {np} previsoes por modelo. Bootstrap B={B:,}, semente {s}."
            ).format(ini=d["inicio"], fim=d["fim"], n=d["n_meses"], cv=d["registros_cv"],
                     jan=b["n_janelas"], h=b["n_horizontes"], np=b["n_previsoes_por_modelo"],
                     B=v["_meta"]["B"], s=v["_meta"]["seed"]).replace(",", ".")


# Arquivos de metrica que alimentam ESTE manuscrito. A lista e explicita de proposito:
# results/ guarda tambem rodadas antigas (amostras sinteticas, e a rodada de 60 meses que
# o paper usa so como comparacao), e essas tem n_predictions diferente por desenho.
# Varrer results/ inteiro faz o verificador acusar coorte misturada onde nao ha, e um
# aviso falso em vermelho no documento e pior que aviso nenhum: ensina a ignorar.
METRICAS_DA_RODADA = [
    "benchmark_sim_real_sp_2010_2023_metrics.csv",
    "benchmark_exog_temp_climatology_metrics.csv",
    "benchmark_exog_temp_observed_metrics.csv",
    "benchmark_slide60_2010_2023_metrics.csv",
]


def coorte_misturada(v: dict) -> list:
    """Avisa se os resultados desta rodada nao forem todos da mesma execucao."""
    import csv
    esperado = v["backtest"]["n_previsoes_por_modelo"]
    fora = []
    for nome in METRICAS_DA_RODADA:
        arq = ROOT / "results" / nome
        if not arq.exists():
            fora.append((nome, "ausente"))
            continue
        with open(arq, encoding="utf-8") as f:
            vals = {int(float(r["n_predictions"])) for r in csv.DictReader(f)
                    if r.get("n_predictions")}
        if vals != {esperado}:
            fora.append((nome, sorted(vals)))
    return fora


def main() -> int:
    BUILD.mkdir(parents=True, exist_ok=True)
    v = json.loads((PAPER / "verified_numbers.json").read_text(encoding="utf-8"))
    tex = (PAPER / "manuscript.tex").read_text(encoding="utf-8")
    cites, refs, ordem = indexa(tex)

    doc = setup_document()
    n_fig = n_tab = n_ausente = 0

    # ---- carimbo de proveniencia ----
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("VERSAO PARA REVISAO, NAO PARA SUBMISSAO")
    r.bold = True
    r.font.size = Pt(10)
    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.0
    r2 = p2.add_run(carimbo(v))
    r2.font.size = Pt(9)
    r2.italic = True

    fora = coorte_misturada(v)
    if fora:
        pw = doc.add_paragraph()
        rw = pw.add_run("[AVISO] coortes misturadas: " + "; ".join(
            f"{a} tem n_predictions={b}" for a, b in fora) +
            ". Nao comparar numeros entre secoes ate resolver.")
        rw.bold = True
        rw.font.color.rgb = VERMELHO
        rw.font.size = Pt(10)

    # ---- titulo e autores ----
    mt = re.search(r"\\title\{(.*?)\n*\}\s*\n\s*\n", tex, re.S)
    if mt:
        titulo = " ".join(x[0] for x in limpa_inline(mt.group(1), cites, refs))
        pt = doc.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = pt.add_run(titulo)
        rt.bold = True
        rt.font.size = Pt(14)
    ma = re.search(r"\\author\{(.*?)\}\s*\n\s*\n", tex, re.S)
    if ma:
        aut = re.sub(r"\\thanks\{.*?\}", "", ma.group(1), flags=re.S)
        aut = re.sub(r"\\footnotemark\[\d\]", "", aut).replace("\\and", ",")
        pa = doc.add_paragraph()
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        escreve_runs(pa, limpa_inline(aut, cites, refs))
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("Faculdade de Saude Publica, Universidade de Sao Paulo, Sao Paulo, Brazil")
    rf.font.size = Pt(10)
    rf.italic = True

    # ---- corpo ----
    corpo = tex[tex.index("\\maketitle") + len("\\maketitle"):]
    corpo = corpo[:corpo.index("\\bibliographystyle")]

    fig_pend = None
    for bloco in re.split(r"\n\s*\n", corpo):
        b = bloco.strip()
        if not b:
            continue

        m = re.match(r"\\(sub)?section\*?\{(.*?)\}", b, re.S)
        if m:
            txt = " ".join(x[0] for x in limpa_inline(m.group(2), cites, refs))
            ph = doc.add_paragraph()
            rh = ph.add_run(txt)
            rh.bold = True
            rh.font.size = Pt(12 if m.group(1) else 14)
            ph.paragraph_format.space_before = Pt(12)
            ph.paragraph_format.space_after = Pt(4)
            b = b[m.end():].strip()
            if not b:
                continue

        mt2 = re.match(r"\\input\{tables/(\w+)\}", b)
        if mt2:
            n_tab += 1
            arq = (PAPER / "tables" / f"{mt2.group(1)}.tex").read_text(encoding="utf-8")
            add_tabela(doc, parse_tabela(arq, cites, refs), f"Table {n_tab}.")
            continue

        if b.startswith("\\begin{figure}"):
            nome = re.search(r"\\input\{figures/(\w+)\}", b)
            cap = re.search(r"\\caption\{(.*?)\}\s*\n\\label", b, re.S)
            n_fig += 1
            legenda = " ".join(x[0] for x in limpa_inline(cap.group(1), cites, refs)) if cap else ""
            if nome:
                if add_figura(doc, nome.group(1), f"Figure {n_fig}.", legenda):
                    pass
                else:
                    n_ausente += 1
            continue

        if b.startswith("\\") and not re.match(r"\\(noindent|textbf|emph|missing|input)", b):
            continue

        runs = limpa_inline(b, cites, refs)
        if not runs:
            continue
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        escreve_runs(pp, runs)

    # ---- referencias ----
    ph = doc.add_paragraph()
    rh = ph.add_run("References")
    rh.bold = True
    rh.font.size = Pt(14)
    ph.paragraph_format.space_before = Pt(14)
    for i, (txt, pend) in enumerate(formata_bib(ordem), 1):
        pr = doc.add_paragraph()
        pr.paragraph_format.line_spacing = 1.0
        pr.paragraph_format.space_after = Pt(6)
        rr = pr.add_run(f"{i}. ")
        rr.bold = True
        rr.font.size = Pt(10)
        r2 = pr.add_run(txt)
        r2.font.size = Pt(10)
        if pend:
            r2.font.color.rgb = VERMELHO
            r2.bold = True

    saida = BUILD / "manuscript_revisao.docx"
    doc.save(saida)
    kb = saida.stat().st_size / 1024
    print(f"  {saida.relative_to(ROOT)} ({kb:.0f} KB)")
    print(f"  {n_fig} figuras embutidas ({n_ausente} ausentes), {n_tab} tabelas nativas")
    print(f"  {len(ordem)} referencias, numeradas por ordem de aparicao")
    print(f"  coorte misturada: {'SIM ' + str(fora) if fora else 'nao'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
